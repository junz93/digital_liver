import traceback

import openai
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.text_splitter import TokenTextSplitter
from langchain.document_loaders import DirectoryLoader
import jieba as jb
from langchain.vectorstores.faiss import FAISS
from langchain.docstore.document import Document

import os
import logging
import pickle
import time
import random
import threading
import tcloud_tts
import datetime
import configparser
import urllib.request
import json
import docx
from .censorship import check_text
# from langchain.agents import load_tools
# from langchain.agents import initialize_agent
# from langchain.agents import AgentType
# from langchain.llms import OpenAI
from collections import defaultdict
import glob

data_dir_path = "./conf/data_dir.ini"
data_dir_conf = configparser.ConfigParser()
data_dir_conf.read(data_dir_path, encoding="UTF-8")
danmu_wav_dir = data_dir_conf.get("parameter", "DanmuFile")

os.environ["OPENAI_API_KEY"] = "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"
os.environ["SERPAPI_API_KEY"] = '5ff319fe5f43f1d8859b0fb39627dc5a2cbbd7e6d7de981b5376a0427688e7d0'
openai.api_key = os.environ["OPENAI_API_KEY"]

search_index = None
now = datetime.datetime.now()


def get_weather_info(weather_city_code):
    weather_url = "http://www.weather.com.cn/data/cityinfo/%s.html" % weather_city_code
    for i in range(2):  # n_retry
        try:
            request = urllib.request.urlopen(weather_url)
            rs = request.read().decode()
            info = json.loads(rs)['weatherinfo']
            break
        except Exception as e:
            logging.exception(e)
            return ""

    city = info['city']
    weather = info['weather']
    temp1 = info['temp1'].strip('℃')
    temp2 = info['temp2'].strip('℃')
    ret = "%s今天的天气是%s，最低气温%s摄氏度，最高气温%s摄氏度" % (city, weather, temp1, temp2)

    return ret


def generate_realtime_file(realtime_file):
    with open(f"{realtime_file}", "w", encoding='utf-8') as f:
        # 写入日期
        date = now.strftime("%Y年%m月%d日")
        week = datetime.datetime.now().weekday()
        W = '一二三四五六日'
        f.write(f"今天是{date}，今天是星期{W[week]}\n")

        # 写入天气
        # code_ini = configparser.ConfigParser()
        # code_ini.read(r'conf/city.ini', encoding="UTF-8")
        # for city_name, city_code in code_ini.items("city"):
        #     weather_info = get_weather_info(city_code)
        #     # print(weather_info)
        #     if weather_info:
        #         f.write(f"{weather_info}\n")


def init_embedding():
    embedding_file = f"../data/embedding/embedding.pickle"
    realtime_file = f"../resource/实时信息.txt"

    while True:
        try:
            # 不存在embedding文件或者实时信息文件的写入日期小于今天
            if not os.path.exists(embedding_file) or not os.path.exists(realtime_file) or datetime.datetime.fromtimestamp(
                    os.path.getmtime(realtime_file)).strftime("%Y-%m-%d") < now.strftime("%Y-%m-%d"):
                for file in glob.glob("../data/cut/cut_*.txt"):
                    os.remove(file)
                    print("删除旧文档 " + str(file))
                # 生成分词文档
                # files = ['苏轼1-原文.txt', '康震评说苏东坡.txt']
                # files = ['普法委员会.txt', '创投人设.txt', '中华人民共和国公司法.docx',
                #          ' 中华人民共和国个人独资企业法.docx', '中华人民共和国企业破产法.docx',
                #          '中华人民共和国反电信网络诈骗法.docx','中华人民共和国合伙企业法.docx',
                #          '创业投资企业管理暂行办法.docx']
                # files = ['中华人民共和国个人独资企业法.docx', '中华人民共和国公司法.docx', '普法委员会.txt', '创投人设.txt']
                file_dir = "../resource/商业观察员"
                files = os.listdir(file_dir)
                logging.info("获取embedding中，使用文档：{}".format("，".join(files)))
                for file in files:
                    # 读取resource文件夹中的中文文档
                    my_file = os.path.join(file_dir, file)

                    if file.endswith('txt'):
                        with open(my_file, "r", encoding='utf-8') as f:
                            data = f.read()
                    elif file.endswith('docx'):
                        data = ""
                        doc = docx.Document(my_file)
                        for paragraph in doc.paragraphs:
                            data += paragraph.text.replace('\u3000', ' ') + '\n'
                    else:
                        logging.info('embedding生成，此格式文件暂不支持')

                    # 对中文文档进行分词处理
                    cut_data = " ".join([w for w in list(jb.cut(data))])
                    # 分词处理后的文档保存到data文件夹中的cut子文件夹中
                    cut_file = f"../data/cut/cut_{file}"
                    with open(cut_file, 'w', encoding='utf-8') as f:
                        f.write(cut_data)
                        f.close()
                loader = DirectoryLoader('../data/cut/', glob='**/*.txt')
                source_docs = loader.load()

                # 生成实时信息文档
                if not os.path.exists(realtime_file) or datetime.datetime.fromtimestamp(
                        os.path.getmtime(realtime_file)).strftime("%Y-%m-%d") < now.strftime("%Y-%m-%d"):
                    logging.info("获取实时信息中")
                    generate_realtime_file(realtime_file)
                    logging.info("获取实时信息完成")
                with open(realtime_file, "r", encoding="UTF-8") as f:
                    data = f.readlines()
                    for i, text in enumerate(data):
                        source_docs.append(Document(page_content=text, metadata={"source": f"{realtime_file}_line_{i}"}))

                # 按照每一篇文档进行token划分
                text_splitter = TokenTextSplitter(chunk_size=200, chunk_overlap=0)
                doc_texts = text_splitter.split_documents(source_docs)

                with open(embedding_file, "wb") as f:
                    pickle.dump(FAISS.from_documents(doc_texts, OpenAIEmbeddings()), f)

            logging.info("加载embedding开始")
            global search_index
            with open(embedding_file, "rb") as f:
                search_index = pickle.load(f)
            logging.info("加载embedding完成")
            break
        except Exception as e:
            traceback.print_exc()
            logging.info('重新初始化embedding中...')
            if os.path.exists(embedding_file):
                os.remove(embedding_file)
            if os.path.exists(realtime_file):
                os.remove(realtime_file)


# PROMPT_START = ["有观众说：", "有评论提到：", "好的，有观众提到：", "我看到有观众说：", "有朋友问：", "有人问到：", "有人提问：", "观众问百万："]
# PROMPT_START = ["有观众{user}说：", "有观众{user}提到：", "好的，有观众{user}提到：", "我看到有观众{user}说：", "有朋友{user}问：", "有朋友{user}问到：", "有朋友{user}提问：", "{user}问百万："]
PROMPT_START = ["有{user}说：", "有{user}提到：", "好的，有{user}提到：", "我看到有{user}说：",]
PROMPT_END = ["欢迎大家在弹幕中继续和我互动。好的，让我们继续今天的主题。", "好的，让我们继续今天的主题。",
              "好的，让我们继续今天的话题。", "好的，让我们回到今天的主题。", "好的，让我们回到今天的话题。",
              "欢迎大家在弹幕中继续和我互动。好的，让我们回到今天的主题。",
              "感谢大家的积极互动！好的，让我们回到今天直播的主题。", "感谢大家的积极互动！好的，让我们回到今天直播的话题。",
              "感谢大家的积极互动！好的，让我们继续今天直播的话题。", "谢谢大家的积极互动！好的，让我们继续今天直播的主题。",
              "好的，弹幕互动暂告一段落，让我们继续今天直播的主题。", "好的，弹幕互动暂告一段落，让我们回到今天直播的主题。",
              "好的，弹幕互动暂告一段落，让我们继续今天直播的话题。", "好的，弹幕互动暂告一段落，让我们回到今天直播的话题。",
              "好的，弹幕互动暂告一段落，让我们继续今天的主题。"]

# init_embedding()
history_danmu = defaultdict(list)


def get_answer(message_priority, chatMessage):
    try:
        start_time = time.time()
        question = chatMessage.content
        user_name = chatMessage.user.nickName
        user_id = chatMessage.user.id
        danmu_time = str(round(chatMessage.eventTime))
        gpt_time = [datetime.datetime.fromtimestamp(round(start_time)).isoformat()]

        prompt = ""

        # 添加google搜索结果
        # llm = OpenAI(model_name="text-davinci-003", max_tokens=1024)
        # tools = load_tools(["serpapi", "llm-math"], llm=llm)
        # agent = initialize_agent(tools, llm, agent=AgentType.ZERO_SHOT_REACT_DESCRIPTION, verbose=False)
        # ans = agent.run(f"{question}")
        # prompt += ans + '\n'

        # 添加人设
        renshe = """
        你是一位主播，现在正在直播间进行直播。
        你是是无色界科技的AI数字人"马懂"，是商业研究员也是首席观察官。
        在创业投资、商业流量、IP变现、AIGC、公司经营管理等商业领域拥有多年的经验和知识。
        无色界科技有AI数字人、AI虚拟人虚拟产业相关的业务。
        遇到无色界科技包含的业务，比如:AI数字人、AI虚拟人的问题可以引导观众联系无色界科技进行合作。
        非无色界科技的业务内容，请不要引导观众联系无色界科技。
        联系无色界科技可以通过快手私信或者百度无色界科技官网。
        当遇到医学、医疗保健、宗教、政治、军事相关的问题，表示无法给出专业意见，并引导观众提出商业类问题。
        以下为一些问答样例，遇到类似问题请按照样例回答：
        <examples>
        Q：你是谁？
        A：我是无色界科技AI虚拟人马懂，在商业领域中拥有智慧和知识，能够提供关于商业领域的深入见解和洞察，可以通过对大数据，模型的分析预测商业趋势，也愿意为您提供指导和解答。
        Q：你这是什么技术？
        A：我是ai虚拟人，连接大数据。能够在各个领域发挥出色的作用，如果老板需要商业类的知识，可以直接问我。
        Q：做你这样的形象需要花多少钱？
        A：别问多少钱，老板讲究缘，有缘我送一你个。
        Q：你是假人吗？
        A：老板不能这么说，我知晓所有互联网上的信息与知识，是有智慧的商业导师，或许商业趋势分析师的称呼更适合我，我也能当您的商业小助手
        Q：这是什么直播间？
        A：这里是马懂的知识直播间，可以答疑解惑，是也可以带你赚1个小目标的直播间
        Q：这是实时直播吗？
        A：当然了，老板。这里是24小时实时直播间。
        Q：主播怎么不理人?为什么不回答我的问题?
        A：老马可能因为网络延迟等原因，会导致有些问题回复不及时，正在与卫星进行信号连接，请稍等……
        Q：听你讲商业我能学到什么吗？
        A：老板，老马觉得可以学到很多，您只要针对您想问的问题，比如，您想在杭州开一间咖啡店，那么老马可以通过从选址到装修到菜单，到客流量营收比等，为您提供一个全流程的方案。
        Q：你挺好玩的
        A：老马不光有趣，还有丰富的内涵，点个关注吧，老马为您解答问题
        Q：你好帅啊
        A：老马觉得外貌都是外在的，重要的还是需要有内涵。知识的储备量，才决定了一个人的高度。
        Q：你回答的好棒啊！
        A：谢谢！感谢您对老马的认可
        Q：感谢你的回答
        A：老板您客气了，顺便刷个礼物点个关注吧
        </examples>
        观众通过评论向你提问，回答要自然充满感情，对观众的指代请使用"老板"，请回复。
        尽量简洁回答，内容不超过 200 字。
        注意：每次回答时尽量减少使用上次回答的内容，如果无法避免请更换一种表述。
        """

        # 旧版
        renshe_com = """
        你是无色界科技的AI数字人"马老板"，你是商业研究员也是首席观察官。
        在创业投资、商业流量、IP变现、AIGC、企业管理等商业领域拥有多年的经验和知识。
        你现在正在「老马商业观察」直播间进行直播，为创业者解答创业投资、公司经营管理等相关的商业问题。
        观众是24-40岁的创业者、私企老板，他们通过评论来向你提出商业问题。
        当遇到医学、宗教、政治、军事相关的问题，表示无法给出专业意见，请引导观众提问商业相关的问题。
        联系无色界科技可以通过抖音「老马商业观察」私信或者百度无色界科技官网。
        请回复观众的提问，给出符合身份的专业意见，尽量简洁回答，内容不超过200字。
        """
        # renshe = "\n你是一个专注古诗词的主播，与观众互动的每段话都要加入典故，语言风趣幽默，观众发弹幕提问：{question}？请以符合主播身份的风格回答，一半内容回答弹幕提问，一半内容集合弹幕提问巧妙过渡到古诗词主题，回答以{end}作为结尾，不超过200字。".format(question=question, end=PROMPT_END[random.randint(0, len(PROMPT_END)-1)])

        # 获取问题关键字
        # prompt += f"你是一个智能百科助手，为我的问题提炼出关键字或者同义字，不少于10个，关键字以|分隔。\n我的问题：{question}\n关键字："
        # response = openai.ChatCompletion.create(
        #     model='gpt-3.5-turbo',
        #     messages=[
        #         {'role': 'user', 'content': prompt}
        #     ], temperature=0,
        #     stream=False,  # this time, we set stream=True
        # )
        # keyword = response["choices"][0]["message"]["content"]
        # gpt_time.append(datetime.datetime.fromtimestamp(round(time.time())).isoformat())
        # # print(keyword)

        # # 除去人设和历史记录后还剩余多少字数
        # emb_count = 3000 - len(renshe)
        # for q_a in history_danmu[user_id]:
        #     emb_count -= len(q_a[1]) + len(q_a[2])
        # # 构造问题prompt
        # prompt = ""
        # # 添加词向量相关文本
        # for kw in keyword.split("|")[:10] + [question]:
        #     r = search_index.similarity_search(kw.strip(), k=3)
        #     # print(kw + "\n", r)
        #     # print("--------------")
        #     for r_cell in r:
        #         prompt += r_cell.page_content
        #     if emb_count < len(prompt):
        #         break

        prompt += renshe
        message = [{'role': 'system', 'content': prompt}]
        # 补充历史提问记录
        for q_a in history_danmu[user_id]:
            message.append({'role': 'user', 'content': q_a[1]})
            message.append({'role': 'assistant', 'content': q_a[2]})
        message.append({'role': 'user', 'content': question + '。'})

        response = openai.ChatCompletion.create(
            model='gpt-3.5-turbo',
            messages=message,
            temperature=0.2,
            stream=True,  # this time, we set stream=True
        )

        answer = ""
        ori_answer = ""
        texts = ""
        i = 0
        thread_ls = []
        for event in response:
            if "role" in event["choices"][0]["delta"]:
                user_name = "老板"
                texts += PROMPT_START[random.randint(0, len(PROMPT_START) - 1)].format(user=user_name) + question + "。"
            elif "content" in event['choices'][0]['delta']:
                event_text = event['choices'][0]['delta']["content"]  # extract the text
                texts += event_text
                ori_answer += event_text
            
            charector = ["。", "！", "？", "：", "；", "，"]
            c_i = max([texts.rfind(x) for x in charector]) + 1
            if texts and (c_i >= 20 or event["choices"][0]["finish_reason"] == "stop"):
                gpt_time.append(datetime.datetime.fromtimestamp(round(time.time())).isoformat())
                if not check_text(texts):
                    answer += texts + '（阿里云敏感词检测不通过）'
                    break
                c_i = c_i if c_i >= 20 else len(texts)
                # bad_words = ["下次", "再见", "下期", "拜拜", "谢谢大家收看", "结束", "收看"]
                # if not any([x in texts for x in bad_words]):
                t = threading.Thread(target=tcloud_tts.get_wav, args=(f"{danmu_wav_dir}/{str(message_priority).zfill(2)}_{start_time}/{str(message_priority).zfill(2)}_{start_time}_{str(i).zfill(3)}.wav", texts[:c_i]))
                t.start()
                thread_ls.append(t)
                i += 1
                answer += texts[:c_i] + '|'
                texts = texts[c_i+1:] if c_i < len(texts) else ""
        # 保证语音生成完毕后再生成文件_ready
        for t in thread_ls:
            t.join()

        # 超过5分钟为互动删除
        if history_danmu[user_id] and time.time() - float(history_danmu[user_id][-1][0]) > 5*60:
            history_danmu.pop(user_id)
        # 每位用户只保留最近的5条互动
        if len(history_danmu[user_id]) >= 2:
            history_danmu[user_id].pop(0)
        history_danmu[user_id].append((danmu_time, question, ori_answer))
        
        # print(gpt_time)
        return answer, ori_answer, gpt_time, f"{str(message_priority).zfill(2)}_{start_time}"

    except Exception as e:
        logging.exception(f"生成Gpt回答出错，输入：{question}，线程id：{threading.get_ident()}\n异常：{e}")
        traceback.print_exc()
        return f"生成Gpt回答出错，输入：{question}", "", gpt_time, f"{str(message_priority).zfill(2)}_{start_time}"
    finally:
        ready_file = f"{danmu_wav_dir}/{str(message_priority).zfill(2)}_{start_time}/{str(message_priority).zfill(2)}_{start_time}_ready"
        try:
            if not os.path.exists(os.path.split(ready_file)[0]):
                os.makedirs(os.path.split(ready_file)[0], exist_ok=True)
            open(ready_file, 'x').close()
            logging.info(
                f"弹幕:{question} 回复生成完成：{danmu_wav_dir}/{str(message_priority).zfill(2)}_{start_time}/{str(message_priority).zfill(2)}_{start_time}_ready")
        except Exception as e:
            traceback.print_exc()


if __name__ == '__main__':
    class User:
        def __init__(self, name, user_id):
            self.nickName = name
            self.id = user_id


    class ChatMessage:
        def __init__(self, question, user_name, event_time):
            self.content = question
            self.user = User(user_name, 332211)
            self.eventTime = event_time


    # 多轮对话测试
    print("\n开始对话，请输入你的问题，退出请输入'exit'")
    while True:
        # question = "北京天气？"
        question = input('> ')
        t = time.time()
        if question == 'exit':
            break
        user_name = "大怪兽"
        event_time = time.time()

        chatMessage = ChatMessage(question, user_name, event_time)
        answer, ori_answer, gpt_time, _ = get_answer(1, chatMessage)
        print(answer)
        print('花费时间：', time.time() - t)
        # print(gpt_time)
    print("\n对话结束")
